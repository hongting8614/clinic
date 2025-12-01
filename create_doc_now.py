# -*- coding: utf-8 -*-
import sys
import os

try:
    from docx import Document
    from docx.shared import Pt, Mm
    from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
    from docx.oxml.ns import qn
    
    # 创建文档
    doc = Document()
    
    # 设置页边距
    for section in doc.sections:
        section.top_margin = Mm(37)
        section.bottom_margin = Mm(35)
        section.left_margin = Mm(28)
        section.right_margin = Mm(26)
    
    # 添加总标题
    para = doc.add_paragraph()
    run = para.add_run('医务室健康演练课目汇总与操作脚本')
    run.font.name = '黑体'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
    run.font.size = Pt(22)
    run.font.bold = True
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    para.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    para.paragraph_format.line_spacing = Pt(28)
    
    # 添加一级标题
    para = doc.add_paragraph()
    run = para.add_run('一、课目一 烧烫伤现场急救演练')
    run.font.name = '黑体'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
    run.font.size = Pt(16)
    run.font.bold = True
    para.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    para.paragraph_format.line_spacing = Pt(28)
    para.paragraph_format.space_before = Pt(12)
    para.paragraph_format.space_after = Pt(6)
    
    # 添加示例正文
    para = doc.add_paragraph()
    run = para.add_run('（一）教学目标')
    run.font.name = '楷体'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '楷体')
    run.font.size = Pt(16)
    para.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    para.paragraph_format.line_spacing = Pt(28)
    para.paragraph_format.space_before = Pt(6)
    para.paragraph_format.space_after = Pt(3)
    
    para = doc.add_paragraph()
    run = para.add_run('1. 了解园区内常见烧烫伤风险场景（餐饮区、热饮、自助取水等）。')
    run.font.name = '仿宋'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋')
    run.font.size = Pt(16)
    para.paragraph_format.first_line_indent = Pt(32)
    para.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    para.paragraph_format.line_spacing = Pt(28)
    
    # 保存文档
    filename = '医务室健康演练课目汇总与操作脚本.docx'
    doc.save(filename)
    
    # 写入成功标记
    with open('success.txt', 'w', encoding='utf-8') as f:
        f.write('文档生成成功！\n')
        f.write(f'文件名：{filename}\n')
        f.write(f'保存路径：{os.path.abspath(filename)}\n')
        
except Exception as e:
    with open('error.txt', 'w', encoding='utf-8') as f:
        f.write(f'错误：{str(e)}\n')
        import traceback
        f.write(traceback.format_exc())
