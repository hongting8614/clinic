#!/usr/bin/env python
# -*- coding: utf-8 -*-
"""
医务室健康演练课目文档生成脚本
运行后会在当前目录生成：医务室健康演练课目汇总与操作脚本.docx

使用方法：
1. 确保已安装 python-docx：pip install python-docx
2. 运行脚本：python generate_doc.py
3. 生成的文档会保存在当前目录
"""

from docx import Document
from docx.shared import Pt, Mm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.ns import qn

def set_font(run, font_name, font_size, bold=False):
    """设置字体（公文格式）"""
    # 字体降级处理
    font_map = {
        '中标宋': '黑体',
        '小标宋': '黑体',
        '仿宋_GB2312': '仿宋',
        '楷体_GB2312': '楷体'
    }
    actual_font = font_map.get(font_name, font_name)
    
    run.font.name = actual_font
    try:
        run._element.rPr.rFonts.set(qn('w:eastAsia'), actual_font)
    except:
        pass
    run.font.size = Pt(font_size)
    run.font.bold = bold

def add_para(doc, text, font_name='仿宋_GB2312', font_size=16, bold=False, align='left', 
             space_before=0, space_after=0, first_line_indent=False, line_spacing_pt=28):
    """添加段落（公文格式）"""
    para = doc.add_paragraph()
    run = para.add_run(text)
    set_font(run, font_name, font_size, bold)
    
    # 对齐方式
    if align == 'center':
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    else:
        para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    
    # 段前段后距离
    if space_before > 0:
        para.paragraph_format.space_before = Pt(space_before)
    if space_after > 0:
        para.paragraph_format.space_after = Pt(space_after)
    
    # 首行缩进（公文正文2字符）
    if first_line_indent:
        para.paragraph_format.first_line_indent = Pt(font_size * 2)  # 2个字符宽度
    
    # 行距固定值（公文格式28磅）
    para.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    para.paragraph_format.line_spacing = Pt(line_spacing_pt)
    
    return para

# 课目数据（可以轻松添加更多）
courses = [
    {
        'title': '一、课目一 烧烫伤现场急救演练',
        'sections': [
            ('（一）教学目标', [
                '1. 了解园区内常见烧烫伤风险场景（餐饮区、热饮、自助取水等）。',
                '2. 掌握"冲、脱、泡、盖、送"的现场处置原则。',
                '3. 能够在无医生在场时，完成基本的冷却、遮盖和求助操作。'
            ]),
            ('（二）角色设置', [
                '1. 医生 1 人。',
                '2. 模拟伤员 1 人（手臂或小腿粘贴假伤口）。',
                '3. 工作人员 1 人（扮演餐饮区服务员）。',
                '4. 学员若干（游客或员工）。'
            ]),
            ('（三）场地与道具', [
                '1. 场地：园区餐饮区或自助取水点附近。',
                '2. 道具：水龙头或水桶、毛巾、剪刀、无菌纱布或干净纱布、烫伤膏（展示用）。'
            ]),
            ('（四）演练步骤', [
                '1. 情景导入',
                '   工作人员口播：模拟游客在取热水时不慎被开水烫到前臂。',
                '2. 发现伤情与报告',
                '   模拟伤员呼喊"烫到了，好疼"。同伴立即协助远离热源，并向工作人员呼救。',
                '3. 医生示范处置流程',
                '   （1）迅速远离热源，确保不再继续受热。',
                '   （2）在流动凉水下持续冲洗伤处 15～20 分钟（演练中可简化为 1～2 分钟）。',
                '   （3）不撕扯粘在烫伤处的衣物，仅剪开周围未粘连部分。',
                '   （4）冲洗后用干净纱布轻轻覆盖伤口，不涂抹牙膏、酱油、酒精等。',
                '   （5）根据面积、部位（面部、关节）及伤情严重程度，判断是否需要立即送医。',
                '4. 学员分组实操',
                '   每 2～3 人一组，轮流扮演伤员和施救者，按步骤完成操作。',
                '5. 医生总结口诀',
                '   再次强调"冲、脱、泡、盖、送"，提醒大家遇到严重烫伤要尽快到医务室或医院。'
            ])
        ]
    },
    {
        'title': '二、课目二 心肺骤停与心肺复苏 / AED 演练',
        'sections': [
            ('（一）教学目标', [
                '1. 识别心肺骤停的主要体征。',
                '2. 掌握基础心肺复苏（CPR）步骤。',
                '3. 会在园区内正确使用 AED 设备。'
            ]),
            ('（二）角色设置', [
                '1. 医生 1 人。',
                '2. CPR 模拟人或仿真人体道具 1 个。',
                '3. 工作人员 1～2 人。',
                '4. 学员若干。'
            ]),
            ('（三）场地与道具', [
                '1. 场地：园区广场、急救培训室或空旷区域。',
                '2. 道具：CPR 模拟人、AED 训练机（如有）、垫子。'
            ]),
            ('（四）演练步骤', [
                '1. 情景导入',
                '   主持人口播：模拟游客在高刺激项目后突然倒地、无反应。',
                '2. 发现与呼救',
                '   （1）发现者轻拍双肩，大声呼喊"你还好吗"，无反应。',
                '   （2）观察 10 秒内胸廓是否起伏，判断无正常呼吸或仅濒死喘息。',
                '   （3）立即大声呼救，指定一人拨打"120"，另一人去取 AED。',
                '3. 医生示范 CPR 流程',
                '   （1）将患者仰卧于坚硬平面上。',
                '   （2）双手交叉，掌根置于胸骨中下段，垂直按压。',
                '   （3）频率 100～120 次/分钟，深度约 5 厘米，按压与回弹时间相等。',
                '   （4）根据培训要求，可演示"30 次按压 + 2 次人工呼吸"的循环。',
                '4. AED 使用演示（如有设备）',
                '   （1）打开电源并按语音提示操作。',
                '   （2）按图示将电极贴于胸部相应位置。',
                '   （3）语音提示分析心律时，所有人远离患者。',
                '   （4）提示需要电击时，按键放电后立即继续胸外按压。',
                '5. 学员轮流实操',
                '   每名学员至少完成 1 分钟标准按压，医生现场逐一纠正手位、力度和节奏。',
                '6. 医生总结口诀',
                '   强调"早识别、早呼救、早按压、早除颤"，并提醒园区关键区域应配备 AED。'
            ])
        ]
    }
    # 在这里可以继续添加其他课目，格式相同
]

def main():
    try:
        doc = Document()
        
        # 设置页边距（公文格式：上37mm 下35mm 左28mm 右26mm）
        sections = doc.sections
        for section in sections:
            section.top_margin = Mm(37)
            section.bottom_margin = Mm(35)
            section.left_margin = Mm(28)
            section.right_margin = Mm(26)
        
        # 总标题（小标宋体二号居中）
        add_para(doc, '医务室健康演练课目汇总与操作脚本', '小标宋', 22, True, 'center', 0, 0)
        
        # 生成所有课目
        for course in courses:
            # 一级标题：课目标题（黑体三号加粗）
            add_para(doc, course['title'], '黑体', 16, True, 'left', 12, 6, False, 28)
            
            # 各个小节
            for section_title, items in course['sections']:
                # 二级标题：小节标题（楷体三号）
                add_para(doc, section_title, '楷体_GB2312', 16, False, 'left', 6, 3, False, 28)
                # 正文内容（仿宋三号，首行缩进2字符，行距28磅）
                for item in items:
                    # 判断是否需要缩进（带序号的正文需要缩进）
                    needs_indent = not item.startswith('   ')  # 如果开头有空格说明是子项，不再缩进
                    add_para(doc, item, '仿宋_GB2312', 16, False, 'left', 0, 0, needs_indent, 28)
        
        # 保存文档
        filename = '医务室健康演练课目汇总与操作脚本.docx'
        doc.save(filename)
        print(f'✓ 文档已生成：{filename}')
        print(f'✓ 当前包含 {len(courses)} 个课目')
        print('✓ 可以打开文档查看效果')
    except Exception as e:
        print(f'✗ 生成失败: {e}')
        import traceback
        traceback.print_exc()

if __name__ == '__main__':
    import sys
    sys.stdout.reconfigure(encoding='utf-8')
    print('开始生成文档...')
    main()
    print('完成！')
